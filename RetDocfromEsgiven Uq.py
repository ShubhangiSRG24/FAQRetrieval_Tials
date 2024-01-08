import pandas as pd
from elasticsearch import Elasticsearch
from docx import Document

# Connect to Elasticsearch
es = Elasticsearch(
    [{'host': 'localhost', 'port': 9200}],
    http_auth=('elastic', 'Hd4hRlSdxroHPN=6_u-d')  # Replace with your credentials
)

# Load user queries from an Excel file
excel_file = 'D://Anaconda3//envs//RetFAQ//LC_UserQuery.xlsx'  # Replace with the path to your Excel file
df = pd.read_excel(excel_file)
queries = df['text'].tolist()  # Replace 'Column_with_queries' with the actual column name

# Name of your Elasticsearch index
index_name = "lcfaq_index"

# Function to search in Elasticsearch and get top 3 answers
def search_es(query):
    response = es.search(index=index_name, body={
        "query": {
            "match": {
                "question": query
            }
        },
        "size": 3  # Retrieve top 3 results
    })

    return [(hit["_source"]["answer"], hit["_id"]) for hit in response['hits']['hits']]

# Retrieve answers and their indexes for each query
results = {query: search_es(query) for query in queries}

# Create a new Word document
doc = Document()

# Add the results to the Word document
for query, answers in results.items():
    doc.add_paragraph(f'Query: {query}\n')
    for idx, (answer, es_index) in enumerate(answers, start=1):
        doc.add_paragraph(f'Answer {idx}: {answer} (Index: {es_index})\n')

# Save the Word document
word_file = 'RetAns_toUquery.docx'  # Replace with your desired output file path
doc.save(word_file)