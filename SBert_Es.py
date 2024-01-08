from sentence_transformers import SentenceTransformer
from elasticsearch import Elasticsearch
import json
import pandas as pd
from docx import Document

# Load pre-trained SBERT model
model = SentenceTransformer('all-mpnet-base-v2')

# Get the size of the embeddings
embedding_size = model.get_sentence_embedding_dimension()
print("The dimension of the model output (embedding size) is:", embedding_size)

#Connect to Elasticsearch
es = Elasticsearch(
    [{'host': 'localhost', 'port': 9200}],
    http_auth=('elastic', 'Hd4hRlSdxroHPN=6_u-d')
)

# Define the index name
index_name = "lcfaq_index"

# Create an Elasticsearch index with the required mapping for embeddings
mapping = {
    "mappings": {
        "properties": {
            "question_embedding": {
                "type": "dense_vector",
                "dims": embedding_size  # Adjust based on your model output dimensions
            },
            # Include other fields as needed
        }
    }
}

# Delete the index if it exists and create a new one
if es.indices.exists(index_name):
    es.indices.delete(index=index_name)
es.indices.create(index=index_name, body=mapping)

# Load and process FAQ data
with open("D://Anaconda3//envs//RetFAQ//output_faqs.json", "r") as file:
    faqs = json.load(file)

for i, faq in enumerate(faqs):
    faq['question_embedding'] = model.encode(faq['question']).tolist()
    es.index(index=index_name, id=i, body=faq)

# Function to perform semantic search
def search_es(query_embedding):
    script_query = {
        "script_score": {
            "query": {"match_all": {}},
            "script": {
                "source": "cosineSimilarity(params.query_vector, 'question_embedding') + 1.0",
                "params": {"query_vector": query_embedding}
            }
        }
    }
    response = es.search(index=index_name, body={"query": script_query, "size": 3})
    return [(hit["_source"]["answer"], hit["_id"]) for hit in response['hits']['hits']]

# Load user queries and generate embeddings
df = pd.read_excel("D://Anaconda3//envs//RetFAQ//LC_UserQuery.xlsx")
queries = df['text'].tolist()
query_embeddings = model.encode(queries)

# Retrieve answers for each query embedding
results = {query: search_es(query_embedding) for query, query_embedding in zip(queries, query_embeddings)}

# Output results to a Word document
doc = Document()
for query, answers in results.items():
    doc.add_paragraph(f'Query: {query}\n')
    for idx, (answer, es_index) in enumerate(answers, start=1):
        doc.add_paragraph(f'Answer {idx}: {answer} (Index: {es_index})\n')

# Save the Word document
doc.save('RetAns_toUquery.docx')
